"""Paper discovery, archiving, and dataset ingestion helpers."""

from __future__ import annotations

import json
import hashlib
import re
import shutil
import tempfile
import zipfile
from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import Any
from urllib import error as urllib_error
from urllib import parse as urllib_parse
from urllib import request as urllib_request
import xml.etree.ElementTree as ET

from backend import db
from fnirs_core.data import SUPPORTED_DATA_SUFFIXES, summarize_file
from fnirs_core.knowledge import extract_text_from_document


ROOT_DIR = Path(__file__).resolve().parent.parent
PAPER_CACHE_DIR = ROOT_DIR / "artifacts" / "papers"
SUPPORTED_PAPER_SUFFIXES = {".pdf", ".docx", ".doc", ".md", ".markdown", ".txt", ".text"}
DATA_URL_SUFFIXES = SUPPORTED_DATA_SUFFIXES - {".zip"} | {".zip"}
USER_AGENT = "fnirs-paper-assistant/1.0 (+local research workflow)"
MAX_DOWNLOAD_BYTES = 120 * 1024 * 1024
MODEL_HINT_LIMIT = 6


class PaperWorkflowError(RuntimeError):
    """Raised when the paper workflow cannot complete."""


@dataclass(slots=True)
class PaperCandidate:
    title: str
    authors: list[str] = field(default_factory=list)
    year: str | None = None
    abstract: str = ""
    source: str = "unknown"
    url: str | None = None
    pdf_url: str | None = None
    doi: str | None = None

    def citation(self) -> str:
        author_text = ", ".join(self.authors[:4]) if self.authors else "Unknown authors"
        suffix = " et al." if len(self.authors) > 4 else ""
        year = f" ({self.year})" if self.year else ""
        return f"{author_text}{suffix}{year}. {self.title}."


@dataclass(slots=True)
class DownloadedFile:
    name: str
    path: str
    url: str | None = None
    suffix: str = ""


@dataclass(slots=True)
class PaperMaterial:
    query: str
    search_query: str
    candidate: PaperCandidate
    workspace: str
    paper_path: str | None = None
    paper_text: str = ""
    data_urls: list[str] = field(default_factory=list)
    data_files: list[DownloadedFile] = field(default_factory=list)
    search_results: list[PaperCandidate] = field(default_factory=list)

    def model_context(self, max_text_chars: int = 18000) -> str:
        payload = {
            "query": self.query,
            "search_query": self.search_query,
            "selected_paper": asdict(self.candidate),
            "paper_path": self.paper_path,
            "data_urls": self.data_urls,
            "data_files": [asdict(item) for item in self.data_files],
            "paper_text_excerpt": self.paper_text[:max_text_chars],
        }
        return json.dumps(payload, ensure_ascii=False, indent=2)


@dataclass(slots=True)
class PaperWorkflowResult:
    title: str
    paper_file: str | None
    report_file: str
    metadata_file: str
    rag_refreshed: bool
    datasets: list[dict[str, Any]] = field(default_factory=list)
    dataset_errors: list[str] = field(default_factory=list)

    def model_context(self) -> str:
        return json.dumps(asdict(self), ensure_ascii=False, indent=2)


def collect_paper_material(query: str) -> PaperMaterial:
    """Find an accessible paper, download original text, and discover direct data files."""

    search_query = _extract_search_query(query)
    candidates = find_paper_candidates(query, search_query)
    candidates = _rank_relevant_candidates(query, candidates)
    if not candidates:
        raise PaperWorkflowError("没有找到可用的论文候选。请提供更具体的题名、关键词、DOI 或 PDF 链接。")

    errors: list[str] = []
    for candidate in candidates:
        workspace = _unique_dir(PAPER_CACHE_DIR / _safe_filename(candidate.title or search_query))
        workspace.mkdir(parents=True, exist_ok=True)
        try:
            paper_path, paper_text, data_files = _download_candidate_assets(candidate, workspace)
            material = PaperMaterial(
                query=query,
                search_query=search_query,
                candidate=candidate,
                workspace=str(workspace),
                paper_path=str(paper_path) if paper_path else None,
                paper_text=paper_text or candidate.abstract,
                data_files=data_files,
                search_results=candidates,
            )
            material.data_urls = discover_data_urls("\n".join([candidate.abstract, material.paper_text]))
            material.data_files.extend(_download_data_urls(material.data_urls, workspace / "data"))
            if material.paper_path or material.paper_text.strip():
                return material
        except Exception as exc:
            errors.append(f"{candidate.title}: {exc}")

    detail = "；".join(errors[:4])
    raise PaperWorkflowError(f"找到了论文候选，但没有可读取的论文原文或摘要。{detail}")


def collect_paper_material_with_model(query: str, llm: Any | None = None) -> PaperMaterial:
    search_query = _extract_search_query(query)
    hints = _model_paper_hints(query, llm) if llm is not None else {}
    candidates = find_paper_candidates(query, search_query, hints=hints)
    candidates = _rank_relevant_candidates(query, candidates)
    if not candidates:
        model_candidates = _model_candidates(hints)
        if model_candidates:
            candidates = _rank_relevant_candidates(query, model_candidates)
    if not candidates:
        raise PaperWorkflowError("没有找到与问题足够相关的论文候选。请提供更具体的题名、关键词、DOI 或 PDF 链接。")

    errors: list[str] = []
    for candidate in candidates:
        workspace = _unique_dir(PAPER_CACHE_DIR / _safe_filename(candidate.title or search_query))
        workspace.mkdir(parents=True, exist_ok=True)
        try:
            paper_path, paper_text, data_files = _download_candidate_assets(candidate, workspace)
            material = PaperMaterial(
                query=query,
                search_query=search_query,
                candidate=candidate,
                workspace=str(workspace),
                paper_path=str(paper_path) if paper_path else None,
                paper_text=paper_text or candidate.abstract,
                data_files=data_files,
                search_results=candidates,
            )
            material.data_urls = discover_data_urls("\n".join([candidate.abstract, material.paper_text]))
            material.data_files.extend(_download_data_urls(material.data_urls, workspace / "data"))
            if material.paper_path or material.paper_text.strip():
                return material
        except Exception as exc:
            errors.append(f"{candidate.title}: {exc}")

    detail = "；".join(errors[:4])
    raise PaperWorkflowError(f"找到了相关论文候选，但没有可读取的论文原文或摘要。{detail}")


def find_paper_candidates(raw_query: str, search_query: str | None = None, hints: dict[str, Any] | None = None) -> list[PaperCandidate]:
    search_query = search_query or _extract_search_query(raw_query)
    candidates: list[PaperCandidate] = []
    candidates.extend(_direct_url_candidates(raw_query))
    candidates.extend(_model_candidates(hints or {}))
    for query in _expanded_search_queries(raw_query, search_query, hints or {}):
        candidates.extend(_semantic_scholar_search(query))
        candidates.extend(_arxiv_search(query))
        candidates.extend(_crossref_search(query))
    return _dedupe_candidates(candidates)[:8]


def finalize_paper_workflow(material: PaperMaterial, report_text: str) -> PaperWorkflowResult:
    from backend import services

    title = material.candidate.title or material.search_query or "paper"
    paper_root = services.EXTRACTED_KNOWLEDGE_DIR / "papers"
    paper_root.mkdir(parents=True, exist_ok=True)
    identity_key = _paper_identity_key(material.candidate)
    target_dir = _find_existing_paper_dir(paper_root, identity_key, title) or paper_root / _paper_folder_name(identity_key, title)
    target_dir.mkdir(parents=True, exist_ok=True)

    paper_file: str | None = None
    if material.paper_path:
        source = Path(material.paper_path)
        if source.exists() and source.suffix.lower() in SUPPORTED_PAPER_SUFFIXES:
            target = target_dir / f"paper{source.suffix.lower()}"
            shutil.copy2(source, target)
            paper_file = str(target)

    report_file = str(_write_docx_report(target_dir / "reading_report.docx", title, report_text, material))
    metadata_file = target_dir / "paper_workflow.json"

    datasets: list[dict[str, Any]] = []
    dataset_errors: list[str] = []
    for data_file in material.data_files:
        try:
            datasets.append(_ingest_dataset_file(Path(data_file.path), data_file.name).model_dump())
        except Exception as exc:
            dataset_errors.append(f"{data_file.name}: {exc}")

    rag_refreshed = False
    try:
        services.refresh_knowledge_base()
        rag_refreshed = True
    except Exception as exc:
        dataset_errors.append(f"RAG 刷新失败：{exc}")

    result = PaperWorkflowResult(
        title=title,
        paper_file=paper_file,
        report_file=report_file,
        metadata_file=str(metadata_file),
        rag_refreshed=rag_refreshed,
        datasets=datasets,
        dataset_errors=dataset_errors,
    )
    metadata_file.write_text(
        json.dumps(
            {
                "paper_identity_key": identity_key,
                "normalized_title": _normalize_identity_text(title),
                "material": {
                    "query": material.query,
                    "search_query": material.search_query,
                    "candidate": asdict(material.candidate),
                    "paper_path": paper_file,
                    "data_urls": material.data_urls,
                    "data_files": [asdict(item) for item in material.data_files],
                },
                "result": asdict(result),
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )
    return result


def discover_data_urls(text: str) -> list[str]:
    urls = []
    for raw_url in re.findall(r"https?://[^\s\]\)）>\"']+", text or ""):
        url = raw_url.rstrip(".,;:")
        suffix = Path(urllib_parse.urlparse(url).path).suffix.lower()
        if suffix in DATA_URL_SUFFIXES:
            urls.append(url)
    return list(dict.fromkeys(urls))


def _model_paper_hints(query: str, llm: Any) -> dict[str, Any]:
    messages = [
        {
            "role": "system",
            "content": (
                "你是科研论文检索助手。请把用户需求改写为英文论文检索信息，只输出 JSON，不要 Markdown。"
                "JSON 字段：search_queries(list[str])、must_include(list[str])、must_not_include(list[str])、papers(list[object])。"
                "papers 中每项包含 title、authors(list[str])、year、doi、url、reason。"
                "如果不能确定具体论文，papers 返回空数组，但 search_queries 必须给出 3-5 个英文检索式。"
            ),
        },
        {"role": "user", "content": query},
    ]
    text = ""
    for piece in llm.stream(messages):
        text += piece
    return _parse_json_object(text)


def _parse_json_object(text: str) -> dict[str, Any]:
    cleaned = text.strip()
    if cleaned.startswith("```"):
        cleaned = re.sub(r"^```(?:json)?", "", cleaned, flags=re.I).strip()
        cleaned = re.sub(r"```$", "", cleaned).strip()
    match = re.search(r"\{.*\}", cleaned, flags=re.S)
    if match:
        cleaned = match.group(0)
    try:
        payload = json.loads(cleaned)
    except Exception:
        return {}
    return payload if isinstance(payload, dict) else {}


def _model_candidates(hints: dict[str, Any]) -> list[PaperCandidate]:
    candidates: list[PaperCandidate] = []
    for item in hints.get("papers") or []:
        if not isinstance(item, dict) or not item.get("title"):
            continue
        candidates.append(
            PaperCandidate(
                title=str(item.get("title") or ""),
                authors=[str(author) for author in item.get("authors") or []],
                year=str(item.get("year")) if item.get("year") else None,
                abstract=str(item.get("reason") or ""),
                source="Ollama candidate",
                url=item.get("url"),
                pdf_url=item.get("pdf_url"),
                doi=item.get("doi"),
            )
        )
    return candidates[:MODEL_HINT_LIMIT]


def _expanded_search_queries(raw_query: str, search_query: str, hints: dict[str, Any]) -> list[str]:
    queries = [search_query]
    for item in hints.get("search_queries") or []:
        if isinstance(item, str) and item.strip():
            queries.append(item.strip())
    for paper in hints.get("papers") or []:
        if isinstance(paper, dict) and paper.get("title"):
            queries.append(str(paper["title"]))
        if isinstance(paper, dict) and paper.get("doi"):
            queries.append(str(paper["doi"]))
    normalized = []
    seen: set[str] = set()
    for query in queries:
        cleaned = re.sub(r"\s+", " ", query).strip()
        key = cleaned.lower()
        if not cleaned or key in seen:
            continue
        seen.add(key)
        normalized.append(cleaned)
    if not normalized:
        normalized.append(raw_query.strip())
    return normalized[:8]


def _rank_relevant_candidates(query: str, candidates: list[PaperCandidate]) -> list[PaperCandidate]:
    if not candidates:
        return []
    query_terms = _content_terms(query)
    ranked: list[tuple[float, PaperCandidate]] = []
    for candidate in _dedupe_candidates(candidates):
        text = " ".join([candidate.title, candidate.abstract, " ".join(candidate.authors), candidate.doi or ""])
        candidate_terms = _content_terms(text)
        score = _candidate_relevance_score(query_terms, candidate_terms, candidate)
        if score >= 0.18 or candidate.source in {"direct_url", "Ollama candidate"}:
            ranked.append((score, candidate))
    ranked.sort(key=lambda item: (item[0], item[1].pdf_url is not None, bool(item[1].abstract)), reverse=True)
    return [candidate for _, candidate in ranked]


def _content_terms(text: str) -> set[str]:
    normalized = text.lower()
    tokens = set(re.findall(r"[a-z0-9][a-z0-9+-]{2,}", normalized))
    chinese_terms = {match for match in re.findall(r"[\u4e00-\u9fff]{2,}", normalized)}
    aliases: set[str] = set()
    if "近红外" in normalized or "fnirs" in normalized or "nirs" in normalized:
        aliases.update({"fnirs", "nirs", "near-infrared", "near", "infrared"})
    if "transformer" in normalized or "注意力" in normalized:
        aliases.update({"transformer", "attention"})
    if "deep learning" in normalized or "深度学习" in normalized:
        aliases.update({"deep", "learning", "neural"})
    return tokens | chinese_terms | aliases


def _candidate_relevance_score(query_terms: set[str], candidate_terms: set[str], candidate: PaperCandidate) -> float:
    if not query_terms:
        return 0.0
    overlap = query_terms & candidate_terms
    score = len(overlap) / max(len(query_terms), 1)
    title_terms = _content_terms(candidate.title)
    score += 0.35 * (len(query_terms & title_terms) / max(len(query_terms), 1))
    if candidate.pdf_url:
        score += 0.05
    if candidate.doi:
        score += 0.03
    return score


def build_paper_report_messages(material: PaperMaterial) -> list[dict[str, str]]:
    return [
        {
            "role": "system",
            "content": (
                "你是严谨的中文科研阅读助手。请基于给定论文材料生成适合保存为 Word 的阅读报告。"
                "不得编造论文没有提供的信息；证据不足处写“原文未说明”或“仅从摘要/节选无法判断”。"
                "报告需要覆盖：题名与引用、研究问题、数据与被试、方法流程、模型/统计方法、主要结果、创新点、局限、"
                "可复现建议、与 fNIRS 深度学习平台的关系、是否发现可入库数据。"
            ),
        },
        {
            "role": "user",
            "content": f"用户请求：{material.query}\n\n论文材料 JSON：\n{material.model_context()}",
        },
    ]


def build_paper_final_messages(material: PaperMaterial, result: PaperWorkflowResult) -> list[dict[str, str]]:
    return [
        {
            "role": "system",
            "content": (
                "你是 fNIRS 平台的中文论文助手。请基于已完成的工作流结果，简洁告知用户论文、Word 阅读报告、RAG 入库、"
                "数据集入库和任何限制。所有路径必须原样给出。不要声称未完成的事情已经完成。"
            ),
        },
        {
            "role": "user",
            "content": (
                f"用户请求：{material.query}\n\n选择的论文：\n{material.model_context(max_text_chars=3000)}"
                f"\n\n工作流结果：\n{result.model_context()}"
            ),
        },
    ]


def build_paper_error_messages(query: str, error_message: str) -> list[dict[str, str]]:
    return [
        {
            "role": "system",
            "content": (
                "你是 fNIRS 平台的中文论文助手。论文工作流失败时，需要说明失败原因，并给出下一步可执行建议，"
                "例如提供 DOI、PDF 链接、更精确题名，或稍后重试。"
            ),
        },
        {"role": "user", "content": f"用户请求：{query}\n\n失败原因：{error_message}"},
    ]


def _download_candidate_assets(candidate: PaperCandidate, workspace: Path) -> tuple[Path | None, str, list[DownloadedFile]]:
    url = candidate.pdf_url or candidate.url
    if not url:
        return None, candidate.abstract, []

    downloaded = _download_url(url, workspace / "source")
    suffix = downloaded.suffix.lower()
    data_files: list[DownloadedFile] = []
    if suffix == ".zip":
        extract_dir = workspace / "archive"
        _safe_extract_zip(downloaded, extract_dir)
        paper_path = _first_supported_file(extract_dir, SUPPORTED_PAPER_SUFFIXES)
        data_files.extend(
            DownloadedFile(name=path.name, path=str(path), suffix=path.suffix.lower())
            for path in _supported_files(extract_dir, SUPPORTED_DATA_SUFFIXES)
        )
    else:
        paper_path = downloaded if suffix in SUPPORTED_PAPER_SUFFIXES else None

    paper_text = ""
    if paper_path:
        try:
            paper_text = extract_text_from_document(paper_path)
        except Exception:
            paper_text = candidate.abstract
    return paper_path, paper_text or candidate.abstract, data_files


def _download_data_urls(urls: list[str], destination: Path) -> list[DownloadedFile]:
    files: list[DownloadedFile] = []
    destination.mkdir(parents=True, exist_ok=True)
    for url in urls[:6]:
        try:
            path = _download_url(url, destination / _safe_filename(Path(urllib_parse.urlparse(url).path).stem or "dataset"))
            if path.suffix.lower() in SUPPORTED_DATA_SUFFIXES:
                files.append(DownloadedFile(name=path.name, path=str(path), url=url, suffix=path.suffix.lower()))
        except Exception:
            continue
    return files


def _direct_url_candidates(query: str) -> list[PaperCandidate]:
    candidates = []
    for url in re.findall(r"https?://[^\s\]\)）>\"']+", query):
        cleaned = url.rstrip(".,;:")
        suffix = Path(urllib_parse.urlparse(cleaned).path).suffix.lower()
        if suffix in SUPPORTED_PAPER_SUFFIXES | {".zip"} or "arxiv.org" in cleaned:
            candidates.append(PaperCandidate(title=Path(urllib_parse.urlparse(cleaned).path).stem or "direct paper", source="direct_url", url=cleaned, pdf_url=_arxiv_pdf_url(cleaned) or cleaned))
    return candidates


def _semantic_scholar_search(query: str) -> list[PaperCandidate]:
    url = (
        "https://api.semanticscholar.org/graph/v1/paper/search?"
        + urllib_parse.urlencode(
            {
                "query": query,
                "limit": "5",
                "fields": "title,abstract,authors,year,venue,url,openAccessPdf,externalIds",
            }
        )
    )
    try:
        payload = _http_json(url)
    except Exception:
        return []
    candidates = []
    for item in payload.get("data", []) if isinstance(payload, dict) else []:
        if not isinstance(item, dict):
            continue
        pdf = item.get("openAccessPdf") or {}
        external_ids = item.get("externalIds") or {}
        candidates.append(
            PaperCandidate(
                title=str(item.get("title") or "untitled paper"),
                authors=[str(author.get("name")) for author in item.get("authors", []) if isinstance(author, dict) and author.get("name")],
                year=str(item.get("year")) if item.get("year") else None,
                abstract=str(item.get("abstract") or ""),
                source="Semantic Scholar",
                url=item.get("url"),
                pdf_url=pdf.get("url") if isinstance(pdf, dict) else None,
                doi=external_ids.get("DOI") if isinstance(external_ids, dict) else None,
            )
        )
    return candidates


def _arxiv_search(query: str) -> list[PaperCandidate]:
    url = (
        "http://export.arxiv.org/api/query?"
        + urllib_parse.urlencode({"search_query": f"all:{query}", "start": "0", "max_results": "5", "sortBy": "relevance"})
    )
    try:
        raw = _http_bytes(url).decode("utf-8", errors="ignore")
        root = ET.fromstring(raw)
    except Exception:
        return []
    ns = {"atom": "http://www.w3.org/2005/Atom"}
    candidates = []
    for entry in root.findall("atom:entry", ns):
        title = " ".join((entry.findtext("atom:title", default="", namespaces=ns) or "").split())
        abstract = " ".join((entry.findtext("atom:summary", default="", namespaces=ns) or "").split())
        year = (entry.findtext("atom:published", default="", namespaces=ns) or "")[:4] or None
        authors = [author.findtext("atom:name", default="", namespaces=ns) for author in entry.findall("atom:author", ns)]
        entry_url = entry.findtext("atom:id", default="", namespaces=ns) or None
        pdf_url = None
        for link in entry.findall("atom:link", ns):
            if link.attrib.get("title") == "pdf" or link.attrib.get("type") == "application/pdf":
                pdf_url = link.attrib.get("href")
                break
        candidates.append(PaperCandidate(title=title or "arXiv paper", authors=[item for item in authors if item], year=year, abstract=abstract, source="arXiv", url=entry_url, pdf_url=pdf_url))
    return candidates


def _crossref_search(query: str) -> list[PaperCandidate]:
    url = "https://api.crossref.org/works?" + urllib_parse.urlencode({"query": query, "rows": "5"})
    try:
        payload = _http_json(url)
    except Exception:
        return []
    items = (((payload or {}).get("message") or {}).get("items") or []) if isinstance(payload, dict) else []
    candidates = []
    for item in items:
        if not isinstance(item, dict):
            continue
        title = " ".join((item.get("title") or ["untitled paper"])[0].split())
        authors = []
        for author in item.get("author") or []:
            if isinstance(author, dict):
                authors.append(" ".join(part for part in [author.get("given"), author.get("family")] if part))
        issued = item.get("issued", {}).get("date-parts", [[None]])
        year = str(issued[0][0]) if issued and issued[0] and issued[0][0] else None
        pdf_url = None
        for link in item.get("link") or []:
            if isinstance(link, dict) and "pdf" in str(link.get("content-type", "")).lower():
                pdf_url = link.get("URL")
                break
        candidates.append(PaperCandidate(title=title, authors=authors, year=year, abstract=str(item.get("abstract") or ""), source="Crossref", url=item.get("URL"), pdf_url=pdf_url, doi=item.get("DOI")))
    return candidates


def _dedupe_candidates(candidates: list[PaperCandidate]) -> list[PaperCandidate]:
    seen: set[str] = set()
    unique: list[PaperCandidate] = []
    for candidate in candidates:
        key = _paper_identity_key(candidate)
        if key in seen:
            continue
        seen.add(key)
        unique.append(candidate)
    unique.sort(key=lambda item: (item.pdf_url is None, not item.abstract, item.source != "direct_url"))
    return unique


def _http_json(url: str) -> Any:
    return json.loads(_http_bytes(url).decode("utf-8", errors="ignore"))


def _http_bytes(url: str) -> bytes:
    request = urllib_request.Request(url, headers={"User-Agent": USER_AGENT, "Accept": "*/*"})
    try:
        with urllib_request.urlopen(request, timeout=30) as response:
            chunks: list[bytes] = []
            total = 0
            while True:
                chunk = response.read(1024 * 1024)
                if not chunk:
                    break
                total += len(chunk)
                if total > MAX_DOWNLOAD_BYTES:
                    raise PaperWorkflowError("下载文件超过大小限制。")
                chunks.append(chunk)
            return b"".join(chunks)
    except urllib_error.URLError as exc:
        raise PaperWorkflowError(f"无法访问论文或数据链接：{exc}") from exc


def _download_url(url: str, destination_base: Path) -> Path:
    data = _http_bytes(url)
    parsed = urllib_parse.urlparse(url)
    suffix = Path(parsed.path).suffix.lower()
    if not suffix:
        suffix = ".pdf" if data[:4] == b"%PDF" else ".txt"
    path = _unique_path(destination_base.with_suffix(suffix))
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_bytes(data)
    return path


def _safe_extract_zip(path: Path, destination: Path) -> None:
    destination.mkdir(parents=True, exist_ok=True)
    root = destination.resolve()
    with zipfile.ZipFile(path) as archive:
        for member in archive.infolist():
            target = (root / member.filename).resolve()
            target.relative_to(root)
            if member.is_dir():
                target.mkdir(parents=True, exist_ok=True)
                continue
            target.parent.mkdir(parents=True, exist_ok=True)
            with archive.open(member) as source, target.open("wb") as output:
                shutil.copyfileobj(source, output)


def _supported_files(root: Path, suffixes: set[str]) -> list[Path]:
    return [
        path
        for path in sorted(root.rglob("*"))
        if path.is_file() and path.suffix.lower() in suffixes and "__MACOSX" not in path.parts and not path.name.startswith(".")
    ]


def _first_supported_file(root: Path, suffixes: set[str]) -> Path | None:
    files = _supported_files(root, suffixes)
    return files[0] if files else None


def _write_docx_report(path: Path, title: str, report_text: str, material: PaperMaterial) -> Path:
    try:
        from docx import Document
    except Exception as exc:  # pragma: no cover - optional dependency
        raise PaperWorkflowError("生成 Word 阅读报告需要 python-docx。") from exc
    document = Document()
    document.add_heading("论文阅读报告", level=0)
    document.add_paragraph(f"论文：{title}")
    document.add_paragraph(f"引用：{material.candidate.citation()}")
    if material.candidate.url:
        document.add_paragraph(f"来源：{material.candidate.url}")
    for line in report_text.replace("\r\n", "\n").split("\n"):
        text = line.strip()
        if not text:
            continue
        if text.startswith("### "):
            document.add_heading(text[4:].strip(), level=3)
        elif text.startswith("## "):
            document.add_heading(text[3:].strip(), level=2)
        elif text.startswith("# "):
            document.add_heading(text[2:].strip(), level=1)
        else:
            document.add_paragraph(text.lstrip("-* "))
    target = path
    document.save(str(target))
    return target


def _ingest_dataset_file(path: Path, filename: str) -> Any:
    from backend import services

    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_DATA_SUFFIXES:
        raise ValueError(f"不支持的数据格式：{suffix}")
    services.DATASET_DIR.mkdir(parents=True, exist_ok=True)
    dataset_id = db.new_id("ds")
    target = services.DATASET_DIR / dataset_id / _safe_filename(filename)
    if not target.suffix:
        target = target.with_suffix(suffix)
    target.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(path, target)
    summary = summarize_file(target)
    timestamp = db.now_iso()
    with db.connect() as connection:
        connection.execute(
            """
            INSERT INTO datasets (id, name, filename, path, suffix, summary_json, created_at, updated_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (dataset_id, target.stem, target.name, str(target), suffix, db.dumps(summary), timestamp, timestamp),
        )
    from backend.schemas import DatasetResponse

    return DatasetResponse(
        id=dataset_id,
        name=target.stem,
        filename=target.name,
        path=str(target),
        suffix=suffix,
        summary=summary,
        created_at=timestamp,
        updated_at=timestamp,
    )


def _extract_search_query(query: str) -> str:
    text = re.sub(r"https?://\S+", " ", query)
    text = re.sub(r"(帮我|请|找到|查找|搜索|论文|阅读报告|word|原文|存入|本地|rag|数据|分析)", " ", text, flags=re.I)
    text = re.sub(r"\s+", " ", text).strip(" ，。,.")
    return text or query.strip()


def _paper_identity_key(candidate: PaperCandidate) -> str:
    if candidate.doi:
        return f"doi_{_safe_filename(candidate.doi.lower())}"
    normalized_title = _normalize_identity_text(candidate.title)
    if normalized_title:
        digest = hashlib.sha1(normalized_title.encode("utf-8")).hexdigest()[:12]
        return f"title_{digest}"
    source = candidate.pdf_url or candidate.url or "paper"
    digest = hashlib.sha1(source.encode("utf-8")).hexdigest()[:12]
    return f"url_{digest}"


def _normalize_identity_text(value: str) -> str:
    text = re.sub(r"[^\w\u4e00-\u9fff]+", " ", (value or "").lower())
    return re.sub(r"\s+", " ", text).strip()


def _paper_folder_name(identity_key: str, title: str) -> str:
    return f"{identity_key}_{_safe_filename(title)[:64]}"


def _find_existing_paper_dir(root: Path, identity_key: str, title: str) -> Path | None:
    normalized_title = _normalize_identity_text(title)
    for candidate in root.glob(f"{identity_key}*"):
        if candidate.is_dir():
            return candidate
    for metadata in root.glob("*/paper_workflow.json"):
        try:
            payload = json.loads(metadata.read_text(encoding="utf-8"))
        except Exception:
            continue
        if payload.get("paper_identity_key") == identity_key:
            return metadata.parent
        if normalized_title and payload.get("normalized_title") == normalized_title:
            return metadata.parent
    return None


def _arxiv_pdf_url(url: str) -> str | None:
    if "arxiv.org/abs/" in url:
        return url.replace("/abs/", "/pdf/") + ("" if url.endswith(".pdf") else ".pdf")
    return None


def _safe_filename(value: str) -> str:
    cleaned = "".join(char if char.isalnum() or char in "-_." else "_" for char in value.strip())
    return cleaned[:100] or "paper"


def _unique_path(path: Path) -> Path:
    if not path.exists():
        return path
    for index in range(1, 1000):
        candidate = path.with_name(f"{path.stem}_{index}{path.suffix}")
        if not candidate.exists():
            return candidate
    raise PaperWorkflowError("无法创建唯一文件名。")


def _unique_dir(path: Path) -> Path:
    if not path.exists():
        return path
    for index in range(1, 1000):
        candidate = path.with_name(f"{path.name}_{index}")
        if not candidate.exists():
            return candidate
    raise PaperWorkflowError("无法创建唯一目录。")

"""Local RAG utilities for the fNIRS self-service platform."""

from __future__ import annotations

import hashlib
import json
import os
import re
import shutil
import subprocess
import tempfile
import unicodedata
import zipfile
from dataclasses import asdict, dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Iterable
from urllib import error as urllib_error
from urllib import request as urllib_request

import numpy as np


SUPPORTED_TEXT_SUFFIXES = {".md", ".markdown", ".txt", ".text", ".pdf", ".docx", ".doc"}
CHINA_TZ = timezone.utc
TEXT_ENCODINGS = ("utf-8-sig", "utf-8", "gb18030", "big5")
UTF16_ENCODINGS = ("utf-16", "utf-16-le", "utf-16-be")


class KnowledgeBaseError(RuntimeError):
    """Raised when the local knowledge base cannot be built or queried."""


def _now() -> str:
    return datetime.now(CHINA_TZ).isoformat()


def _is_cjk(char: str) -> bool:
    return (
        "\u3400" <= char <= "\u4dbf"
        or "\u4e00" <= char <= "\u9fff"
        or "\uf900" <= char <= "\ufaff"
    )


def _text_quality_score(text: str) -> float:
    if not text:
        return 0.0

    replacement_count = text.count("\ufffd")
    control_count = sum(
        1
        for char in text
        if unicodedata.category(char)[0] == "C" and char not in "\n\r\t"
    )
    cjk_count = sum(1 for char in text if _is_cjk(char))
    mojibake_count = sum(text.count(marker) for marker in ("Ã", "Â", "â", "¤", "¥", "§", "¨", "©"))
    printable_count = sum(1 for char in text if char.isprintable() or char in "\n\r\t")

    return (
        printable_count
        + cjk_count * 2.0
        - replacement_count * 20.0
        - control_count * 12.0
        - mojibake_count * 3.0
    )


def _repair_utf8_mojibake(text: str) -> str:
    best = text
    best_score = _text_quality_score(text)
    for encoding in ("latin-1", "cp1252"):
        try:
            candidate = text.encode(encoding).decode("utf-8")
        except UnicodeError:
            continue
        score = _text_quality_score(candidate)
        if score > best_score + 2 and any(_is_cjk(char) for char in candidate):
            best = candidate
            best_score = score
    return best


def _decode_text_bytes(raw: bytes) -> str:
    if not raw:
        return ""

    candidates: list[tuple[float, int, str]] = []
    encodings = list(TEXT_ENCODINGS)
    if raw.startswith((b"\xff\xfe", b"\xfe\xff")) or raw.count(b"\x00") / len(raw) > 0.1:
        encodings.extend(UTF16_ENCODINGS)

    for order, encoding in enumerate(encodings):
        try:
            text = raw.decode(encoding)
        except UnicodeError:
            continue
        text = _repair_utf8_mojibake(text)
        candidates.append((_text_quality_score(text), -order, text))

    if candidates:
        return max(candidates)[2]
    return raw.decode("utf-8", errors="replace")


def _glyph_name_token_ratio(text: str) -> float:
    tokens = re.findall(r"/G[0-9A-Fa-f]{2,4}", text)
    stripped = re.sub(r"\s+", "", text)
    if not stripped:
        return 0.0
    return sum(len(token) for token in tokens) / len(stripped)


def _looks_like_pdf_glyph_names(text: str) -> bool:
    if not text.strip():
        return False
    tokens = re.findall(r"/G[0-9A-Fa-f]{2,4}", text)
    if len(tokens) < 10:
        return False
    cjk_count = sum(1 for char in text if _is_cjk(char))
    return _glyph_name_token_ratio(text) > 0.35 and cjk_count < max(len(tokens) // 10, 3)


def _looks_like_readable_pdf_text(text: str) -> bool:
    if not text.strip():
        return False
    if _looks_like_pdf_glyph_names(text):
        return False
    printable_count = sum(1 for char in text if char.isprintable() or char in "\n\r\t")
    return printable_count >= max(len(text) * 0.8, 1)


def _extract_pdf_with_pypdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except Exception as exc:  # pragma: no cover - optional dependency
        raise KnowledgeBaseError("PDF parsing requires the `pypdf` package.") from exc

    reader = PdfReader(str(path))
    pages: list[str] = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    return "\n".join(pages)


def _extract_pdf_with_pdfplumber(path: Path) -> str | None:
    try:
        import pdfplumber  # type: ignore[import-not-found]
    except Exception:
        return None

    pages: list[str] = []
    try:
        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                pages.append(page.extract_text() or "")
    except Exception:
        return None
    return "\n".join(pages)


def _extract_pdf_with_pdftotext(path: Path) -> str | None:
    executable = shutil.which("pdftotext")
    if not executable:
        return None

    with tempfile.TemporaryDirectory() as tmpdir:
        output_path = Path(tmpdir) / "output.txt"
        command = [
            executable,
            "-enc",
            "UTF-8",
            "-layout",
            str(path),
            str(output_path),
        ]
        try:
            subprocess.run(command, check=True, capture_output=True, timeout=60)
        except Exception:
            return None
        if not output_path.exists():
            return None
        return _decode_text_bytes(output_path.read_bytes())


def _extract_pdf_text(path: Path) -> str:
    primary_text = _extract_pdf_with_pypdf(path)
    if _looks_like_readable_pdf_text(primary_text):
        return primary_text

    candidates = [primary_text] if primary_text and primary_text.strip() else []
    for extractor in (_extract_pdf_with_pdftotext, _extract_pdf_with_pdfplumber):
        text = extractor(path)
        if text and text.strip():
            if _looks_like_readable_pdf_text(text):
                return text
            candidates.append(text)

    readable = [text for text in candidates if _looks_like_readable_pdf_text(text)]
    if readable:
        return max(readable, key=_text_quality_score)

    if candidates and any(_looks_like_pdf_glyph_names(text) for text in candidates):
        raise KnowledgeBaseError(
            "PDF text extraction produced unreadable font glyph codes. "
            "Please upload a PDF with an embedded text Unicode map, or convert it to DOCX/TXT before upload."
        )

    return max(candidates, key=_text_quality_score) if candidates else ""


def _read_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".md", ".markdown", ".txt", ".text"}:
        return _decode_text_bytes(path.read_bytes())

    if suffix == ".pdf":
        return _extract_pdf_text(path)

    if suffix == ".docx":
        try:
            from docx import Document
        except Exception as exc:  # pragma: no cover - optional dependency
            raise KnowledgeBaseError("DOCX parsing requires the `python-docx` package.") from exc

        document = Document(str(path))
        lines = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        return "\n".join(lines)

    if suffix == ".doc":
        raw = path.read_bytes()
        if raw.startswith(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"):
            raise KnowledgeBaseError("Legacy binary .doc files are not supported. Please save the document as .docx or .txt.")
        return _decode_text_bytes(raw)

    raise KnowledgeBaseError(f"Unsupported file suffix: {suffix}")


def extract_text_from_document(path: Path) -> str:
    """Extract text from a supported source document.

    Zip archives are flattened and the first supported file is extracted.
    """

    path = Path(path)
    if not path.exists():
        raise KnowledgeBaseError(f"Document does not exist: {path}")

    suffix = path.suffix.lower()
    if suffix == ".zip":
        with tempfile.TemporaryDirectory() as tmpdir:
            with zipfile.ZipFile(path) as archive:
                _safe_extract_zip(archive, Path(tmpdir))
            candidates = [
                candidate
                for candidate in Path(tmpdir).rglob("*")
                if candidate.is_file() and candidate.suffix.lower() in SUPPORTED_TEXT_SUFFIXES - {".zip"}
            ]
            if not candidates:
                raise KnowledgeBaseError("Zip archive does not contain a supported document.")
            return extract_text_from_document(candidates[0])

    return _read_text(path)


def _safe_extract_zip(archive: zipfile.ZipFile, destination: Path) -> None:
    destination = destination.resolve()
    for member in archive.infolist():
        target = (destination / member.filename).resolve()
        try:
            target.relative_to(destination)
        except ValueError as exc:
            raise KnowledgeBaseError("Zip archive contains an unsafe path.") from exc
        if member.is_dir():
            target.mkdir(parents=True, exist_ok=True)
            continue
        target.parent.mkdir(parents=True, exist_ok=True)
        with archive.open(member) as source, target.open("wb") as output:
            output.write(source.read())


def document_id_from_source(source: str) -> str:
    return hashlib.sha1(source.encode("utf-8")).hexdigest()[:16]


def _normalized_text(text: str) -> str:
    return re.sub(r"\s+", " ", text.lower()).strip()


def _chunk_text(text: str, *, chunk_size: int, chunk_overlap: int) -> list[str]:
    if not text.strip():
        return []

    paragraphs = [part.strip() for part in re.split(r"\n\s*\n", text) if part.strip()]
    if not paragraphs:
        paragraphs = [text.strip()]

    chunks: list[str] = []
    current = ""
    for paragraph in paragraphs:
        candidate = f"{current}\n\n{paragraph}".strip() if current else paragraph
        if len(candidate) <= chunk_size:
            current = candidate
            continue
        if current:
            chunks.append(current.strip())
        if len(paragraph) <= chunk_size:
            current = paragraph
        else:
            start = 0
            overlap = min(max(chunk_overlap, 0), max(chunk_size - 1, 0))
            while start < len(paragraph):
                end = min(start + chunk_size, len(paragraph))
                chunks.append(paragraph[start:end].strip())
                if end >= len(paragraph):
                    current = ""
                    break
                start = max(end - overlap, start + 1)
    if current.strip():
        chunks.append(current.strip())
    return [chunk for chunk in chunks if chunk]


def _hash_embedding(texts: list[str], dimension: int) -> np.ndarray:
    dimension = max(int(dimension), 64)
    vectors = np.zeros((len(texts), dimension), dtype=np.float32)
    for row, text in enumerate(texts):
        normalized = _normalized_text(text)
        if not normalized:
            continue
        tokens = re.findall(r"[\w\u4e00-\u9fff]+", normalized) or normalized.split()
        grams: list[str] = []
        for token in tokens:
            if len(token) <= 2:
                grams.append(token)
                continue
            grams.extend(token[i : i + 3] for i in range(max(len(token) - 2, 1)))
        for gram in grams or tokens:
            digest = hashlib.blake2b(gram.encode("utf-8"), digest_size=8).digest()
            index = int.from_bytes(digest, "little") % dimension
            vectors[row, index] += 1.0
        norm = float(np.linalg.norm(vectors[row]))
        if norm:
            vectors[row] /= norm
    return vectors


@dataclass(slots=True)
class DocumentChunk:
    chunk_id: str
    source: str
    title: str
    content: str
    order: int
    enabled: bool = True


@dataclass(slots=True)
class KnowledgeDocument:
    id: str
    source: str
    title: str
    path: str
    suffix: str
    size_chars: int
    chunk_count: int
    updated_at: str | None
    managed: bool

    def to_summary(self) -> dict[str, Any]:
        return asdict(self)


@dataclass(slots=True)
class SearchResult:
    source: str
    title: str
    content: str
    score: float
    order: int

    def to_dict(self) -> dict[str, Any]:
        snippet = self.content.strip().replace("\r", "").replace("\n", " ")
        return {
            "source": self.source,
            "title": self.title,
            "score": round(float(self.score), 4),
            "order": self.order,
            "snippet": snippet[:320],
        }


@dataclass(slots=True)
class KnowledgeStats:
    total_documents: int
    total_chunks: int
    source_files: list[str]
    source_roots: list[str]
    vector_store_path: str
    embedding_model: str
    embedding_dim: int
    index_updated_at: str | None = None

    def to_dict(self) -> dict[str, Any]:
        return asdict(self)


class OllamaEmbeddingClient:
    """Tiny Ollama embedding client using the native local API."""

    def __init__(self, *, model: str, base_url: str, timeout: float = 600.0, batch_size: int = 8) -> None:
        self.model = model
        self.base_url = self._normalize_base_url(base_url)
        self.timeout = timeout
        self.batch_size = max(int(batch_size), 1)

    @staticmethod
    def _normalize_base_url(base_url: str) -> str:
        return base_url.rstrip("/")

    def embed(self, texts: list[str]) -> np.ndarray:
        if not texts:
            return np.empty((0, 0), dtype=np.float32)

        batches: list[np.ndarray] = []
        for start in range(0, len(texts), self.batch_size):
            batches.append(self._embed_batch(texts[start : start + self.batch_size]))
        return np.vstack(batches).astype(np.float32)

    def _embed_batch(self, texts: list[str]) -> np.ndarray:
        payload = {"model": self.model, "input": texts}
        request = urllib_request.Request(
            f"{self.base_url}/api/embed",
            data=json.dumps(payload, ensure_ascii=False).encode("utf-8"),
            headers={"Content-Type": "application/json"},
            method="POST",
        )
        try:
            with urllib_request.urlopen(request, timeout=self.timeout) as response:
                data = json.loads(response.read().decode("utf-8"))
        except urllib_error.URLError as exc:  # pragma: no cover - network path
            raise KnowledgeBaseError(f"Cannot connect to Ollama embedding service: {exc.reason}") from exc
        except Exception as exc:  # pragma: no cover - network path
            raise KnowledgeBaseError(f"Ollama embedding response is invalid: {exc}") from exc

        embeddings = data.get("embeddings")
        if embeddings is None and "embedding" in data:
            embeddings = [data["embedding"]]
        if not isinstance(embeddings, list) or len(embeddings) != len(texts):
            raise KnowledgeBaseError("Ollama embedding response did not match the input batch.")
        return np.asarray(embeddings, dtype=np.float32)


class KnowledgeBase:
    """A small persistent vector store for local fNIRS documents."""

    fallback_embedding_model = "local-hashing-char-ngram"

    def __init__(
        self,
        sources: Iterable[Path | str],
        *,
        chunk_size: int = 1400,
        chunk_overlap: int = 180,
        vector_store_dir: Path | str | None = None,
        embedding_model: str | None = None,
        embedding_base_url: str | None = None,
        embedding_dim: int = 4096,
        managed_roots: Iterable[Path | str] | None = None,
        allow_embedding_fallback: bool = True,
    ) -> None:
        self.source_roots = [Path(source).resolve() for source in sources]
        self.managed_roots = [Path(root).resolve() for root in (managed_roots or [])]
        self.chunk_size = max(int(chunk_size), 200)
        self.chunk_overlap = max(int(chunk_overlap), 0)
        self.vector_store_dir = Path(vector_store_dir or "artifacts/vector_store").resolve()
        self.vector_store_dir.mkdir(parents=True, exist_ok=True)
        self.vectors_path = self.vector_store_dir / "vectors.npz"
        self.metadata_path = self.vector_store_dir / "metadata.json"
        self.embedding_model = embedding_model or os.getenv("FNIRS_EMBEDDING_MODEL", "qwen3-embedding:8b")
        self.embedding_base_url = embedding_base_url or os.getenv("FNIRS_OLLAMA_BASE_URL", "http://localhost:11434")
        self.embedding_dim = int(embedding_dim)
        self.allow_embedding_fallback = allow_embedding_fallback
        self.documents: list[DocumentChunk] = []
        self._source_index: dict[str, Path] = {}
        self._source_size_chars: dict[str, int] = {}
        self._configured_source_roots = list(self.source_roots)
        self._vectors: np.ndarray | None = None
        self._index_updated_at: str | None = None
        self._active_embedding_model = self.embedding_model

    def refresh(self) -> None:
        files = self._discover_files()
        if not files:
            raise KnowledgeBaseError("No knowledge files were found for the RAG index.")

        previous = self._load_metadata()
        enabled_map: dict[str, bool] = {}
        if previous:
            for document in previous.get("documents", []):
                for chunk in document.get("chunks", []):
                    enabled_map[str(chunk.get("chunk_id"))] = bool(chunk.get("enabled", True))

        chunks: list[DocumentChunk] = []
        source_index: dict[str, Path] = {}
        source_size_chars: dict[str, int] = {}
        for path in files:
            relative_source = self._to_relative_source(path)
            raw_text = self._read_text_file(path)
            if _looks_like_pdf_glyph_names(raw_text):
                continue
            if not raw_text.strip():
                continue
            source_index[relative_source] = path
            source_size_chars[relative_source] = len(raw_text)
            for order, content in enumerate(_chunk_text(raw_text, chunk_size=self.chunk_size, chunk_overlap=self.chunk_overlap)):
                chunk_id = f"{relative_source}#{order}"
                chunks.append(
                    DocumentChunk(
                        chunk_id=chunk_id,
                        source=relative_source,
                        title=path.stem,
                        content=content,
                        order=order,
                        enabled=enabled_map.get(chunk_id, True),
                    )
                )

        if not chunks:
            raise KnowledgeBaseError("Knowledge files were found, but no readable text chunks were produced.")

        self.documents = chunks
        self._source_index = source_index
        self._source_size_chars = source_size_chars
        self._vectors = self._embed_texts([chunk.content for chunk in chunks])
        self.embedding_dim = int(self._vectors.shape[1]) if self._vectors.ndim == 2 and self._vectors.size else self.embedding_dim
        self._index_updated_at = _now()
        self._persist_vector_store()

    def add_or_update_document(self, path: Path | str) -> None:
        path = Path(path).resolve()
        if path.suffix.lower() not in SUPPORTED_TEXT_SUFFIXES:
            raise KnowledgeBaseError(f"Unsupported file suffix: {path.suffix.lower()}")
        if not path.exists():
            raise KnowledgeBaseError(f"Document does not exist: {path}")

        loaded_existing = self._load_for_incremental_update(path)
        if not loaded_existing:
            self.refresh()
            return

        raw_text = self._read_text_file(path)
        if _looks_like_pdf_glyph_names(raw_text) or not raw_text.strip():
            raise KnowledgeBaseError("Knowledge document did not produce readable text chunks.")

        source = self._to_relative_source(path)
        previous_enabled = {
            chunk.order: chunk.enabled
            for chunk in self.documents
            if chunk.source == source
        }
        new_chunks = [
            DocumentChunk(
                chunk_id=f"{source}#{order}",
                source=source,
                title=path.stem,
                content=content,
                order=order,
                enabled=previous_enabled.get(order, True),
            )
            for order, content in enumerate(
                _chunk_text(raw_text, chunk_size=self.chunk_size, chunk_overlap=self.chunk_overlap)
            )
        ]
        if not new_chunks:
            raise KnowledgeBaseError("Knowledge document did not produce readable text chunks.")

        remaining_indices = [index for index, chunk in enumerate(self.documents) if chunk.source != source]
        existing_vectors = self._vectors
        if existing_vectors is not None and len(remaining_indices) and existing_vectors.shape[0] == len(self.documents):
            remaining_vectors = existing_vectors[remaining_indices]
        else:
            remaining_vectors = None

        remaining_documents = [self.documents[index] for index in remaining_indices]
        new_vectors = self._embed_texts([chunk.content for chunk in new_chunks])
        if remaining_vectors is None:
            vectors = new_vectors
        else:
            vectors = np.vstack([remaining_vectors, new_vectors]).astype(np.float32)

        self.documents = remaining_documents + new_chunks
        self._source_index = {
            source_name: source_path
            for source_name, source_path in self._source_index.items()
            if source_name != source
        }
        self._source_index[source] = path
        self._source_size_chars = {
            source_name: size_chars
            for source_name, size_chars in self._source_size_chars.items()
            if source_name != source
        }
        self._source_size_chars[source] = len(raw_text)
        self._vectors = vectors
        self.embedding_dim = int(vectors.shape[1]) if vectors.ndim == 2 and vectors.size else self.embedding_dim
        self._index_updated_at = _now()
        self._persist_vector_store()

    def _load_for_incremental_update(self, target_path: Path) -> bool:
        if self.documents and self._vectors is not None:
            return True
        if not (self.vectors_path.exists() and self.metadata_path.exists()):
            return False

        metadata = self._load_metadata() or {}
        target_source = self._to_relative_source(target_path)
        metadata_without_target = dict(metadata)
        metadata_without_target["documents"] = [
            document
            for document in metadata.get("documents", [])
            if str(document.get("source", "")) != target_source
        ]
        if not self._metadata_matches_sources(metadata_without_target, extra_sources={target_source}):
            return False

        self._index_updated_at = metadata.get("index_updated_at")
        self._active_embedding_model = metadata.get("embedding_model", self.embedding_model)
        self.embedding_dim = int(metadata.get("embedding_dim", self.embedding_dim))
        stored = np.load(self.vectors_path, allow_pickle=True)
        self._vectors = np.asarray(stored["vectors"], dtype=np.float32)
        self._load_documents_from_metadata(metadata)
        if self._vectors.shape[0] != len(self.documents):
            return False
        return True

    def stats(self) -> KnowledgeStats:
        self._ensure_loaded()
        documents = self.list_documents()
        source_files = [document.source for document in documents]
        return KnowledgeStats(
            total_documents=len(documents),
            total_chunks=len(self.documents),
            source_files=source_files,
            source_roots=[str(root) for root in self.source_roots],
            vector_store_path=str(self.vector_store_dir),
            embedding_model=self._active_embedding_model,
            embedding_dim=int(self.embedding_dim),
            index_updated_at=self._index_updated_at,
        )

    def list_sources(self, *, limit: int = 50) -> list[str]:
        self._ensure_loaded()
        items = sorted(self._source_index)
        return items[: max(int(limit), 1)]

    def list_documents(self) -> list[KnowledgeDocument]:
        self._ensure_loaded()
        chunk_counts: dict[str, int] = {}
        for chunk in self.documents:
            chunk_counts[chunk.source] = chunk_counts.get(chunk.source, 0) + 1

        documents: list[KnowledgeDocument] = []
        for source, path in sorted(self._source_index.items()):
            stat = path.stat() if path.exists() else None
            documents.append(
                KnowledgeDocument(
                    id=document_id_from_source(source),
                    source=source,
                    title=path.stem,
                    path=str(path),
                    suffix=path.suffix.lower(),
                    size_chars=self._source_size_chars.get(source, 0),
                    chunk_count=chunk_counts.get(source, 0),
                    updated_at=datetime.fromtimestamp(stat.st_mtime, timezone.utc).isoformat() if stat else None,
                    managed=self._is_managed_path(path),
                )
            )
        return documents

    def find_document(self, document_id: str) -> KnowledgeDocument | None:
        for document in self.list_documents():
            if document.id == document_id:
                return document
        return None

    def get_document(self, document_id: str) -> tuple[KnowledgeDocument, str]:
        document = self.find_document(document_id)
        if document is None:
            raise KnowledgeBaseError(f"Unknown knowledge document id: {document_id}")
        return document, self._read_text_file(Path(document.path))

    def get_document_chunks(self, document_id: str) -> list[DocumentChunk]:
        document = self.find_document(document_id)
        if document is None:
            raise KnowledgeBaseError(f"Unknown knowledge document id: {document_id}")
        return [chunk for chunk in self.documents if chunk.source == document.source]

    def set_chunk_enabled(self, document_id: str, order: int, enabled: bool) -> DocumentChunk:
        document = self.find_document(document_id)
        if document is None:
            raise KnowledgeBaseError(f"Unknown knowledge document id: {document_id}")
        for chunk in self.documents:
            if chunk.source == document.source and chunk.order == order:
                chunk.enabled = bool(enabled)
                self._persist_vector_store()
                return chunk
        raise KnowledgeBaseError(f"Unknown chunk order {order} for document {document_id}.")

    def get_source_text(self, source: str, max_chars: int = 5000) -> str:
        self._ensure_loaded()
        path = self._source_index.get(source)
        if path is None:
            for candidate_source, candidate_path in self._source_index.items():
                if candidate_source.endswith(source):
                    path = candidate_path
                    break
        if path is None:
            raise KnowledgeBaseError(f"Unknown source: {source}")
        text = self._read_text_file(path)
        return text[: max(int(max_chars), 1)]

    def search(self, query: str, *, top_k: int = 4) -> list[SearchResult]:
        if not query or not query.strip():
            return []
        self._ensure_loaded()
        if self._vectors is None or not len(self.documents):
            return []

        query_vector = self._embed_texts([query.strip()])[0]
        if query_vector.shape[0] != self._vectors.shape[1]:
            self.refresh()
            query_vector = self._embed_texts([query.strip()])[0]
        if len(self.documents) != int(self._vectors.shape[0]):
            self.refresh()
            query_vector = self._embed_texts([query.strip()])[0]
        scores = np.asarray(self._vectors @ query_vector, dtype=np.float32).ravel()
        if scores.size == 0:
            return []

        top_k = max(int(top_k), 1)
        ranked_indices = np.argsort(scores)[::-1]
        results: list[SearchResult] = []
        for index in ranked_indices:
            chunk = self.documents[int(index)]
            if not chunk.enabled:
                continue
            score = float(scores[int(index)])
            if score <= 0 and results:
                break
            results.append(
                SearchResult(
                    source=chunk.source,
                    title=chunk.title,
                    content=chunk.content,
                    score=score,
                    order=chunk.order,
                )
            )
            if len(results) >= top_k:
                break
        return results

    def _discover_files(self) -> list[Path]:
        files: list[Path] = []
        for root in self.source_roots:
            if not root.exists():
                continue
            for path in sorted(root.rglob("*")):
                if path.is_file() and path.suffix.lower() in SUPPORTED_TEXT_SUFFIXES:
                    files.append(path.resolve())
        return files

    def _discover_indexable_files(self) -> list[Path]:
        files: list[Path] = []
        for path in self._discover_files():
            if path.suffix.lower() in {".md", ".markdown", ".txt", ".text"}:
                try:
                    if _looks_like_pdf_glyph_names(_decode_text_bytes(path.read_bytes())):
                        continue
                except Exception:
                    continue
            files.append(path)
        return files

    def _to_relative_source(self, path: Path) -> str:
        path = path.resolve()
        for root in self.source_roots:
            try:
                relative = path.relative_to(root.resolve())
                return relative.as_posix()
            except ValueError:
                continue
        return path.name

    def _is_managed_path(self, path: Path) -> bool:
        resolved = path.resolve()
        for root in self.managed_roots:
            try:
                resolved.relative_to(root.resolve())
                return True
            except ValueError:
                continue
        return False

    def _read_text_file(self, path: Path) -> str:
        return extract_text_from_document(path)

    def _embed_texts(self, texts: list[str]) -> np.ndarray:
        if not texts:
            return np.empty((0, self.embedding_dim), dtype=np.float32)

        if self.embedding_model == self.fallback_embedding_model:
            self._active_embedding_model = self.fallback_embedding_model
            return _hash_embedding(texts, self.embedding_dim)

        try:
            client = OllamaEmbeddingClient(
                model=self.embedding_model,
                base_url=self.embedding_base_url,
            )
            vectors = client.embed(texts)
            if vectors.size:
                self._active_embedding_model = self.embedding_model
                return self._normalize(vectors)
        except Exception:
            if not self.allow_embedding_fallback:
                raise

        self._active_embedding_model = self.fallback_embedding_model
        return _hash_embedding(texts, self.embedding_dim)

    def _persist_vector_store(self) -> None:
        self.vector_store_dir.mkdir(parents=True, exist_ok=True)
        if self._vectors is None:
            raise KnowledgeBaseError("Vector store is empty.")

        documents_payload: list[dict[str, Any]] = []
        for source, path in sorted(self._source_index.items()):
            source_chunks = [item for item in self.documents if item.source == source]
            if not source_chunks:
                continue
            stat = path.stat() if path.exists() else None
            documents_payload.append(
                {
                    "id": document_id_from_source(source),
                    "source": source,
                    "title": path.stem,
                    "path": str(path),
                    "suffix": path.suffix.lower(),
                    "size_chars": self._source_size_chars.get(source, 0),
                    "chunk_count": len(source_chunks),
                    "updated_at": datetime.fromtimestamp(stat.st_mtime, timezone.utc).isoformat() if stat else None,
                    "managed": self._is_managed_path(path),
                    "chunks": [
                        {
                            "chunk_id": item.chunk_id,
                            "order": item.order,
                            "enabled": item.enabled,
                            "source": item.source,
                            "title": item.title,
                        }
                        for item in source_chunks
                    ],
                }
            )

        np.savez_compressed(
            self.vectors_path,
            vectors=self._vectors.astype(np.float32),
            active_embedding_model=np.asarray([self._active_embedding_model]),
            configured_embedding_model=np.asarray([self.embedding_model]),
            embedding_dim=np.asarray([self.embedding_dim], dtype=np.int64),
        )
        payload = {
            "index_updated_at": self._index_updated_at,
            "configured_embedding_model": self.embedding_model,
            "embedding_base_url": self.embedding_base_url.rstrip("/"),
            "embedding_model": self._active_embedding_model,
            "embedding_dim": int(self.embedding_dim),
            "source_roots": [str(root) for root in self.source_roots],
            "documents": documents_payload,
        }
        self.metadata_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")

    def _load_metadata(self) -> dict[str, Any] | None:
        if not self.metadata_path.exists():
            return None
        try:
            return json.loads(self.metadata_path.read_text(encoding="utf-8"))
        except Exception:
            return None

    def _ensure_loaded(self) -> None:
        if self.documents and self._vectors is not None:
            return
        if self.vectors_path.exists() and self.metadata_path.exists():
            metadata = self._load_metadata() or {}
            if not self._metadata_matches_sources(metadata):
                self.refresh()
                return
            self._index_updated_at = metadata.get("index_updated_at")
            self._active_embedding_model = metadata.get("embedding_model", self.embedding_model)
            self.embedding_dim = int(metadata.get("embedding_dim", self.embedding_dim))
            stored = np.load(self.vectors_path, allow_pickle=True)
            self._vectors = np.asarray(stored["vectors"], dtype=np.float32)
            self._load_documents_from_metadata(metadata)
            if self._vectors.shape[0] != len(self.documents):
                self.refresh()
            return
        self.refresh()

    def _load_documents_from_metadata(self, metadata: dict[str, Any]) -> None:
        self._source_index = {}
        self._source_size_chars = {}
        self.documents = []
        enabled_map: dict[str, bool] = {}
        for document in metadata.get("documents", []):
            source = str(document.get("source", ""))
            path = Path(document.get("path", "")).resolve()
            for chunk in document.get("chunks", []):
                enabled_map[str(chunk.get("chunk_id", ""))] = bool(chunk.get("enabled", True))
            if path.exists():
                raw_text = self._read_text_file(path)
                if _looks_like_pdf_glyph_names(raw_text):
                    continue
                self._source_index[source] = path
                self._source_size_chars[source] = len(raw_text)
                for order, content in enumerate(
                    _chunk_text(raw_text, chunk_size=self.chunk_size, chunk_overlap=self.chunk_overlap)
                ):
                    chunk_id = f"{source}#{order}"
                    self.documents.append(
                        DocumentChunk(
                            chunk_id=chunk_id,
                            source=source,
                            title=str(document.get("title", path.stem)),
                            content=content,
                            order=order,
                            enabled=enabled_map.get(chunk_id, True),
                        )
                    )

    def _metadata_matches_sources(self, metadata: dict[str, Any], extra_sources: set[str] | None = None) -> bool:
        configured_embedding_model = metadata.get("configured_embedding_model")
        if configured_embedding_model is not None and str(configured_embedding_model) != self.embedding_model:
            return False

        stored_base_url = metadata.get("embedding_base_url")
        if stored_base_url is not None and str(stored_base_url).rstrip("/") != self.embedding_base_url.rstrip("/"):
            return False

        active_embedding_model = metadata.get("embedding_model")
        if configured_embedding_model is None and active_embedding_model not in {None, self.fallback_embedding_model, self.embedding_model}:
            return False

        metadata_roots = {str(Path(root).resolve()) for root in metadata.get("source_roots", [])}
        current_roots = {str(root.resolve()) for root in self._configured_source_roots}
        if metadata_roots and metadata_roots != current_roots:
            return False

        current_sources = {self._to_relative_source(path): path for path in self._discover_indexable_files()}
        metadata_documents = metadata.get("documents", [])
        metadata_sources = {str(document.get("source", "")) for document in metadata_documents}
        allowed_sources = set(metadata_sources)
        allowed_sources.update(extra_sources or set())
        if set(current_sources) != allowed_sources:
            return False

        for document in metadata_documents:
            source = str(document.get("source", ""))
            path = current_sources.get(source)
            if path is None or not path.exists():
                return False
            stat = path.stat()
            updated_at = datetime.fromtimestamp(stat.st_mtime, timezone.utc).isoformat()
            if document.get("updated_at") != updated_at:
                return False
            try:
                if int(document.get("size_chars", -1)) != len(self._read_text_file(path)):
                    return False
            except Exception:
                return False
        return True

    @staticmethod
    def _normalize(vectors: np.ndarray) -> np.ndarray:
        norms = np.linalg.norm(vectors, axis=1, keepdims=True)
        norms[norms == 0] = 1.0
        return (vectors / norms).astype(np.float32)


def build_default_knowledge_base(
    refresh: bool = False,
    *,
    embedding_model: str | None = None,
    embedding_base_url: str | None = None,
) -> KnowledgeBase:
    root = Path(__file__).resolve().parent.parent
    base_knowledge = root / "knowledge" / "base"
    uploaded_knowledge = root / "knowledge" / "uploads" / "extracted"
    vector_store = root / "artifacts" / "vector_store"
    kb = KnowledgeBase(
        [base_knowledge, uploaded_knowledge],
        vector_store_dir=vector_store,
        embedding_model=embedding_model,
        embedding_base_url=embedding_base_url,
        managed_roots=[uploaded_knowledge],
    )
    if refresh:
        kb.refresh()
    else:
        try:
            kb._ensure_loaded()
        except Exception:
            kb.refresh()
    return kb
